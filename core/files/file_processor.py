"""
JARVIS 1.5 — File Processor
SYSTEMS™ · architectofscale.com

Verarbeite JEDE Art von Datei — PDFs, Bilder, Spreadsheets, Code, Vertraege.

Features:
  - PDF: Text extrahieren, zusammenfassen, Vertragsanalyse
  - Bilder: OCR, Beschreibung, Logo-Erkennung
  - Spreadsheets: CSV/XLSX parsen, Analyse, Visualisierung
  - Code: Review, Security Check, Dokumentation
  - Dokumente: DOCX, TXT — Zusammenfassung, Key Points
  - Audio: Transkription (via Voice Engine)

Jede Datei wird automatisch:
  1. Erkannt (MIME Type)
  2. Extrahiert (Text/Daten)
  3. Analysiert (via LLM)
  4. Im Wissen gespeichert (Knowledge Base)
"""

import os
import io
import json
import logging
import mimetypes
from pathlib import Path
from typing import Optional

logger = logging.getLogger("jarvis.files")


class FileProcessor:
    """
    Universeller Datei-Prozessor.
    Kann jede Datei analysieren und verarbeiten.
    """

    def __init__(self, llm_client=None, knowledge_base=None, db_client=None):
        self.llm = llm_client
        self.knowledge = knowledge_base
        self.db = db_client

    async def process(
        self,
        file_data: bytes,
        filename: str,
        user_instruction: str = "",
        auto_store: bool = True,
    ) -> dict:
        """
        Verarbeite eine Datei.
        Erkennt automatisch den Typ und waehlt die passende Pipeline.
        """
        file_type = self._detect_type(filename)
        file_size = len(file_data)

        logger.info(f"Processing file: {filename} ({file_type}, {file_size} bytes)")

        # Text extrahieren
        extracted = await self._extract(file_data, filename, file_type)

        if not extracted.get("text") and not extracted.get("data"):
            return {
                "filename": filename,
                "type": file_type,
                "error": "Konnte keine Inhalte extrahieren",
            }

        # LLM-Analyse
        analysis = None
        if self.llm:
            analysis = await self._analyze(extracted, filename, file_type, user_instruction)

        # Im Wissen speichern
        if auto_store and self.knowledge and analysis:
            await self._store_knowledge(filename, file_type, analysis)

        return {
            "filename": filename,
            "type": file_type,
            "size_bytes": file_size,
            "extracted": extracted,
            "analysis": analysis,
            "stored": auto_store,
        }

    def _detect_type(self, filename: str) -> str:
        """Erkenne den Dateityp."""
        ext = Path(filename).suffix.lower()
        type_map = {
            ".pdf": "pdf",
            ".png": "image", ".jpg": "image", ".jpeg": "image",
            ".gif": "image", ".webp": "image", ".bmp": "image",
            ".csv": "spreadsheet", ".xlsx": "spreadsheet",
            ".xls": "spreadsheet", ".tsv": "spreadsheet",
            ".docx": "document", ".doc": "document",
            ".txt": "text", ".md": "text", ".rtf": "text",
            ".py": "code", ".js": "code", ".ts": "code",
            ".java": "code", ".go": "code", ".rs": "code",
            ".html": "code", ".css": "code", ".sql": "code",
            ".json": "data", ".yaml": "data", ".yml": "data",
            ".xml": "data", ".toml": "data",
            ".mp3": "audio", ".wav": "audio", ".ogg": "audio",
            ".m4a": "audio", ".flac": "audio",
        }
        return type_map.get(ext, "unknown")

    # ═══════════════════════════════════════════════════
    # EXTRACTION
    # ═══════════════════════════════════════════════════

    async def _extract(self, data: bytes, filename: str, file_type: str) -> dict:
        """Extrahiere Inhalte basierend auf Dateityp."""
        extractors = {
            "pdf": self._extract_pdf,
            "image": self._extract_image,
            "spreadsheet": self._extract_spreadsheet,
            "document": self._extract_document,
            "text": self._extract_text,
            "code": self._extract_text,
            "data": self._extract_text,
        }

        extractor = extractors.get(file_type)
        if not extractor:
            return {"text": "", "error": f"Unsupported type: {file_type}"}

        try:
            return await extractor(data, filename)
        except Exception as e:
            logger.error(f"Extraction failed for {filename}: {e}")
            return {"text": "", "error": str(e)}

    async def _extract_pdf(self, data: bytes, filename: str) -> dict:
        """PDF Text extrahieren."""
        text = ""
        pages = 0

        # Versuch 1: PyPDF2 / pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            pages = len(reader.pages)
            text = "\n\n".join(
                page.extract_text() or "" for page in reader.pages
            )
            return {"text": text, "pages": pages, "method": "pypdf"}
        except ImportError:
            pass

        # Versuch 2: pdfminer
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            text = pdfminer_extract(io.BytesIO(data))
            return {"text": text, "method": "pdfminer"}
        except ImportError:
            pass

        # Versuch 3: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                pages = len(pdf.pages)
                text = "\n\n".join(page.extract_text() or "" for page in pdf.pages)
            return {"text": text, "pages": pages, "method": "pdfplumber"}
        except ImportError:
            pass

        return {"text": "", "error": "No PDF library installed (pip install pypdf)"}

    async def _extract_image(self, data: bytes, filename: str) -> dict:
        """Bild analysieren (OCR + Beschreibung)."""
        result = {"text": "", "description": ""}

        # OCR mit pytesseract
        try:
            from PIL import Image
            import pytesseract

            img = Image.open(io.BytesIO(data))
            ocr_text = pytesseract.image_to_string(img, lang="deu+eng")
            result["text"] = ocr_text.strip()
            result["size"] = f"{img.width}x{img.height}"
            result["method"] = "tesseract"
        except ImportError:
            result["ocr_note"] = "pytesseract not installed"

        # Bild-Beschreibung via LLM (Vision)
        if self.llm:
            try:
                import base64
                b64 = base64.b64encode(data).decode()
                ext = Path(filename).suffix.lower().lstrip(".")
                mime = f"image/{ext}" if ext != "jpg" else "image/jpeg"
                result["description"] = f"[Image: {filename}, base64 available for vision model]"
            except Exception:
                pass

        return result

    async def _extract_spreadsheet(self, data: bytes, filename: str) -> dict:
        """Spreadsheet parsen (CSV, XLSX)."""
        ext = Path(filename).suffix.lower()

        # CSV
        if ext in (".csv", ".tsv"):
            import csv
            delimiter = "\t" if ext == ".tsv" else ","
            text = data.decode("utf-8", errors="replace")
            reader = csv.reader(io.StringIO(text), delimiter=delimiter)
            rows = list(reader)
            return {
                "text": text[:5000],
                "rows": len(rows),
                "columns": len(rows[0]) if rows else 0,
                "headers": rows[0] if rows else [],
                "sample": rows[:5] if rows else [],
                "method": "csv",
            }

        # XLSX
        if ext in (".xlsx", ".xls"):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True)
                sheets = {}
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    rows = []
                    for row in ws.iter_rows(max_row=100, values_only=True):
                        rows.append([str(c) if c is not None else "" for c in row])
                    sheets[sheet_name] = {
                        "rows": len(rows),
                        "columns": len(rows[0]) if rows else 0,
                        "headers": rows[0] if rows else [],
                        "sample": rows[:5],
                    }
                return {
                    "text": json.dumps(sheets, indent=2, ensure_ascii=False)[:5000],
                    "sheets": sheets,
                    "sheet_count": len(sheets),
                    "method": "openpyxl",
                }
            except ImportError:
                return {"text": "", "error": "openpyxl not installed"}

        return {"text": "", "error": f"Unsupported spreadsheet: {ext}"}

    async def _extract_document(self, data: bytes, filename: str) -> dict:
        """DOCX Text extrahieren."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return {"text": text, "paragraphs": len(doc.paragraphs), "method": "python-docx"}
        except ImportError:
            return {"text": "", "error": "python-docx not installed"}

    async def _extract_text(self, data: bytes, filename: str) -> dict:
        """Plaintext / Code / JSON lesen."""
        try:
            text = data.decode("utf-8", errors="replace")
            return {
                "text": text,
                "lines": text.count("\n") + 1,
                "chars": len(text),
            }
        except Exception as e:
            return {"text": "", "error": str(e)}

    # ═══════════════════════════════════════════════════
    # ANALYSIS
    # ═══════════════════════════════════════════════════

    async def _analyze(
        self,
        extracted: dict,
        filename: str,
        file_type: str,
        user_instruction: str,
    ) -> dict:
        """Analysiere extrahierte Inhalte mit LLM."""
        text = extracted.get("text", "")[:6000]  # LLM-Limit

        if not text:
            return {"summary": "Keine Textinhalte gefunden."}

        # Analyse-Prompt basierend auf Dateityp
        prompts = {
            "pdf": (
                f"Analysiere dieses PDF-Dokument ({filename}):\n\n"
                f"{text}\n\n"
                f"Erstelle:\n"
                f"1. Zusammenfassung (3-5 Saetze)\n"
                f"2. Key Points (Bullet List)\n"
                f"3. Dokumenttyp (Vertrag, Report, Brief, etc.)\n"
                f"4. Wichtige Daten/Zahlen\n"
                f"5. Action Items (falls vorhanden)"
            ),
            "spreadsheet": (
                f"Analysiere diese Tabelle ({filename}):\n\n"
                f"Headers: {extracted.get('headers', [])}\n"
                f"Rows: {extracted.get('rows', 0)}\n"
                f"Sample:\n{text[:2000]}\n\n"
                f"Erstelle:\n"
                f"1. Was zeigen die Daten?\n"
                f"2. Trends oder Muster\n"
                f"3. Auffaelligkeiten\n"
                f"4. Empfohlene Visualisierungen"
            ),
            "code": (
                f"Reviewe diesen Code ({filename}):\n\n"
                f"```\n{text[:3000]}\n```\n\n"
                f"Erstelle:\n"
                f"1. Was macht der Code?\n"
                f"2. Code-Qualitaet (1-10)\n"
                f"3. Security Issues\n"
                f"4. Verbesserungsvorschlaege"
            ),
            "image": (
                f"Bild analysiert: {filename}\n"
                f"OCR Text: {text[:1000]}\n\n"
                f"Beschreibe was im Bild/Text zu sehen ist."
            ),
        }

        prompt = prompts.get(file_type, (
            f"Analysiere diese Datei ({filename}, Typ: {file_type}):\n\n"
            f"{text[:3000]}\n\n"
            f"Erstelle eine Zusammenfassung und Key Points."
        ))

        # User-Instruction anhaengen
        if user_instruction:
            prompt += f"\n\nZusaetzliche Anweisung: {user_instruction}"

        try:
            analysis_text = await self.llm.generate(
                prompt=prompt,
                model="tier2-llama",
                max_tokens=800,
            )

            return {
                "summary": analysis_text,
                "file_type": file_type,
                "content_length": len(text),
            }
        except Exception as e:
            return {"summary": f"Analyse fehlgeschlagen: {e}"}

    # ═══════════════════════════════════════════════════
    # KNOWLEDGE STORAGE
    # ═══════════════════════════════════════════════════

    async def _store_knowledge(self, filename: str, file_type: str, analysis: dict):
        """Speichere Analyse-Ergebnisse in der Knowledge Base."""
        try:
            summary = analysis.get("summary", "")
            if summary and len(summary) > 20:
                await self.knowledge.store_fact(
                    subject=f"Datei: {filename}",
                    fact=summary[:500],
                    source=f"file_processor/{file_type}",
                    confidence=0.8,
                )
                logger.info(f"File knowledge stored: {filename}")
        except Exception as e:
            logger.debug(f"Could not store file knowledge: {e}")

    # ═══════════════════════════════════════════════════
    # SPECIALIZED ANALYSIS
    # ═══════════════════════════════════════════════════

    async def analyze_contract(self, file_data: bytes, filename: str) -> dict:
        """Spezielle Vertragsanalyse — Donna's Spezialitaet."""
        extracted = await self._extract(file_data, filename, "pdf")
        text = extracted.get("text", "")

        if not text or not self.llm:
            return {"error": "Kein Text oder LLM nicht verfuegbar"}

        analysis = await self.llm.generate(
            prompt=(
                f"Du bist ein Vertrags-Analyst. Analysiere diesen Vertrag:\n\n"
                f"{text[:6000]}\n\n"
                f"Erstelle eine strukturierte Analyse:\n"
                f"1. VERTRAGSTYP: (Dienstleistung, Kauf, Miete, NDA, etc.)\n"
                f"2. PARTEIEN: Wer sind die Vertragspartner?\n"
                f"3. LAUFZEIT: Start, Ende, Kuendigungsfristen\n"
                f"4. WICHTIGE KLAUSELN: Die 5 wichtigsten Punkte\n"
                f"5. RISIKEN: Potenzielle Probleme oder unfaire Klauseln\n"
                f"6. KOSTEN: Alle finanziellen Verpflichtungen\n"
                f"7. EMPFEHLUNG: Unterschreiben oder nachverhandeln?"
            ),
            model="tier1-sonnet" if len(text) > 3000 else "tier2-llama",
            max_tokens=1200,
        )

        return {
            "filename": filename,
            "type": "contract_analysis",
            "analysis": analysis,
            "pages": extracted.get("pages", 0),
        }

    async def analyze_invoice(self, file_data: bytes, filename: str) -> dict:
        """Rechnungsanalyse — Satoshi's Domaene."""
        extracted = await self._extract(file_data, filename, "pdf")
        text = extracted.get("text", "")

        if not text or not self.llm:
            return {"error": "Kein Text oder LLM nicht verfuegbar"}

        analysis = await self.llm.generate(
            prompt=(
                f"Analysiere diese Rechnung:\n\n{text[:4000]}\n\n"
                f"Extrahiere:\n"
                f"1. Rechnungsnummer\n"
                f"2. Datum\n"
                f"3. Absender (Firma, Adresse)\n"
                f"4. Positionen (Beschreibung, Menge, Preis)\n"
                f"5. Netto, MwSt, Brutto\n"
                f"6. Zahlungsziel\n"
                f"7. Bankverbindung\n"
                f"Antworte als strukturiertes JSON."
            ),
            model="tier2-llama",
            max_tokens=600,
        )

        return {
            "filename": filename,
            "type": "invoice_analysis",
            "analysis": analysis,
        }
